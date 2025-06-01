import sqlite3
import logging
from pathlib import Path
from docx import Document
from datetime import datetime

# Настройка логирования
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('bot.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

class CertificateGenerator:
    def __init__(self):
        self.db = sqlite3.connect('certificates.db', check_same_thread=False, timeout=30)
        self.db.execute("PRAGMA journal_mode=WAL")
        self.create_tables()
        self.template_paths = {
            'Курс по электробезопасности': {
                'expert': Path('Элбез Expert Solution.docx'),
                'regular': Path('Элбез 2.docx')
            },
            'Безопасность и Охрана труда': Path('templates2.docx'),
            'Пожарная безопасность в объеме пожарно технического минимума': {
                'Для ответственных': Path('Птм ответсвенный.docx'),
                'Для работников': Path('Птм рабочие.docx')
            },
            'Промышленная безопасность РК по производству работ на опасных производственных объектах': Path(
                'templates3.docx')
        }
        self.output_dir = Path('generated_docs')
        self.output_dir.mkdir(exist_ok=True)

    def create_tables(self):
        self.db.execute('''CREATE TABLE IF NOT EXISTS certificates (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            protocol_number TEXT,
            fullname TEXT,
            workplace TEXT,
            job_title TEXT,
            position TEXT,
            group_number TEXT,
            cert_date DATE,
            next_date DATE,
            template_type TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )''')
        self.db.commit()

    def generate_document(self, data):
        try:
            generated_files = []
            protocol_number = str(self.db.execute("SELECT COUNT(*) FROM certificates").fetchone()[0] + 1).zfill(3)

            if data['position'] == 'Курс по электробезопасности':
                # Генерируем оба сертификата для электробезопасности
                group = data['qualification_group']
                if isinstance(group, str):
                    group = int(group) if group.isdigit() else 0
                replacements = {
                    '{{ input_fullname }}': data['fullname'],
                    '{{ input_place }}': data['workplace'],
                    '{{ input_jobtitle }}': data['job_title'],
                    '{{ input_what }}': "и выше" if group in [4, 5] else "",
                    '{{ input_do }}': "до " if group in [4, 5] else "",
                    '{{ input_nextdate }}': data['next_date'].strftime('%d.%m.%Y') if data['next_date'] else "",
                    '{{ input_proto }}': protocol_number,
                    '{{ input_date }}': data['cert_date'].strftime('%d.%m.%Y'),
                    '{{ input_group }}': str(group)
                }

                # Первый сертификат - Expert Solution
                doc1 = Document(self.template_paths['Курс по электробезопасности']['expert'])
                self._replace_text_in_document(doc1, replacements)
                output_path1 = self._save_document(doc1, f"{data['fullname']}_Expert")
                generated_files.append(output_path1)

                # Второй сертификат - Элбез 2
                doc2 = Document(self.template_paths['Курс по электробезопасности']['regular'])
                self._replace_text_in_document(doc2, replacements)
                output_path2 = self._save_document(doc2, f"{data['fullname']}_2")
                generated_files.append(output_path2)

            else:
                # Определяем путь к шаблону для других курсов
                if data['position'] == 'Пожарная безопасность в объеме пожарно технического минимума':
                    template_path = self.template_paths[data['position']][data['template_type']]
                else:
                    template_path = self.template_paths[data['position']]

                # Формируем замены в зависимости от типа курса
                if data['position'] == 'Безопасность и Охрана труда':
                    replacements = {
                        '{{ input_data }}': data['cert_date'].strftime('%d.%m.%Y'),
                        '{{ input_fullname }}': data['fullname'],
                        '{{ input_place }}': data['workplace'],
                        '{{ input_job }}': data['job_title'],
                        '{{ input_check }}': data['qualification_group'],
                        '{{ input_proto }}': protocol_number
                    }
                elif data[
                    'position'] == 'Промышленная безопасность РК по производству работ на опасных производственных объектах':
                    replacements = {
                        '{{ input_date }}': data['cert_date'].strftime('%d.%m.%Y'),
                        '{{ input_name }}': data['fullname'],
                        '{{ input_place }}': data['workplace'],
                        '{{ input_job }}': data['job_title'],
                        '{{ input_proto }}': protocol_number
                    }
                else:  # Пожарная безопасность
                    replacements = {
                        '{{ input_name }}': data['fullname'],
                        '{{ input_place }}': data['workplace'],
                        '{{ input_jobtitle }}': data['job_title'],
                        '{{ input_nextdate }}': data['next_date'].strftime('%d.%m.%Y') if data['next_date'] else "",
                        '{{ input_proto }}': protocol_number,
                        '{{ input_date }}': data['cert_date'].strftime('%d.%m.%Y')
                    }

                doc = Document(template_path)
                self._replace_text_in_document(doc, replacements)
                output_path = self._save_document(doc, data['fullname'])
                generated_files.append(output_path)

            # Сохраняем в базу данных
            self._save_to_database(data, protocol_number)
            return True, generated_files

        except Exception as e:
            logger.error(f"Ошибка при генерации документа: {str(e)}")
            return False, None

    def _replace_text_in_document(self, doc, replacements):
        # Замена в параграфах
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, str(value))

        # Замена в таблицах
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in replacements.items():
                        if key in cell.text:
                            cell.text = cell.text.replace(key, str(value))

    def _save_document(self, doc, fullname):
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_path = self.output_dir / f"{fullname}_{timestamp}.docx"
        doc.save(output_path)
        return output_path

    def _save_to_database(self, data, protocol_number):
        try:
            self.db.execute('''INSERT INTO certificates 
                (protocol_number, fullname, workplace, position, group_number, 
                 cert_date, next_date, template_type, job_title)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)''',
                            (protocol_number,
                             data['fullname'],
                             data['workplace'],
                             data['position'],
                             str(data['qualification_group']),
                             data['cert_date'],
                             data['next_date'],
                             data.get('template_type', ''),
                             data['job_title']))
            self.db.commit()
        except Exception as e:
            logger.error(f"Ошибка сохранения в БД: {e}")
            # Продолжаем работу даже при ошибке БД
            pass 