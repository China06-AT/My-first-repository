from telegram import Update, ReplyKeyboardRemove, ReplyKeyboardMarkup
from telegram.ext import Application, CommandHandler, MessageHandler, filters, ContextTypes, ConversationHandler
from datetime import datetime, timedelta
import os, logging, sys, sqlite3
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ROW_HEIGHT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import parse_xml

# Состояния диалога
(PEOPLE_COUNT, FULLNAME, WORKPLACE, JOB_TITLE, POSITION,
 GROUP, CHECK_TYPE, CERT_DATE, NEXT_DATE, CONFIRM_DATA, CONFIRM_RESTART) = range(11)

# Настройка логирования
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('bot.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Доступные курсы
POSITIONS = [
    'Курс по электробезопасности',
    'Безопасность и Охрана труда',
    'Пожарная безопасность',
    'Промышленная безопасность'
]

# Опции для пожарной безопасности
FIRE_SAFETY_OPTIONS = [
    '1. Для ответственных',
    '2. Для работников'
]

# Варианты команды старт
start_variations = [
    '/start', 'начать', 'старт', 'запуск', 'приступить', 'перейти к делу',
    'начать работу', 'начинаем', 'готов к началу', 'поехали', 'стартуем',
    'приступаем', 'давайте начнём', 'запустить', 'готов', 'начинаю',
    'перейти к задаче', 'перейти к сути', 'активировать', 'начало работы',
    'начнём', 'давайте приступим', 'хочу начать', 'давайте начинать',
    'готов к запуску', 'стартуем!', 'готов к работе', 'давайте начинать работу',
    'го', 'давай', 'создать', 'начало', 'старт!', 'поехали!', 'начинаем!',
    'СТАРТ', 'СТАРТУЕМ', 'СТАРТУЙТЕ', 'СТАРТУЙ'
]

# Шаблоны сообщений для многократного использования
MESSAGES = {
    'start': "👋 Здравствуйте!\n\nВведите количество человек для которых нужно сделать сертификат:",
    'invalid_count': "❌ Пожалуйста, введите положительное целое число",
    'enter_names': "👤 Введите ФИО для {} человек:\n(Введите каждое ФИО с новой строки)",
    'invalid_names': "❌ Следующие ФИО введены неверно:\n{}\n\nВведите все {} ФИО заново, каждое с новой строки",
    'enter_workplaces': "🏢 Введите названия организаций для {} человек:\n(Каждое с новой строки)",
    'enter_positions': "💼 Введите должности для {} человек:\n(Каждая с новой строки)",
    'choose_course': "📚 Выберите курс обучения для {}:",
    'enter_date': "📅 Введите дату выдачи удостоверения (дд.мм.гггг):\nПример: 01.02.2024",
    'invalid_date': "❌ Неверный формат даты.\n\nВведите дату в формате дд.мм.гггг\nПример: 01.02.2024",
    'future_date': "❌ Дата не может быть в будущем.\nВведите дату в формате дд.мм.гггг:",
    'confirm_data': "📋 Проверьте введенные данные:\n\n{}\nСоздать документы?",
    'generating': "⚙️ Начинаю генерацию документов...",
    'success': "✅ Готово! Для создания новых документов нажмите /start",
    'error': "❌ Произошла ошибка. Попробуйте снова: /start",
    'cancelled': "❌ Операция отменена. Для начала нажмите /start"
}

# Обновляем словарь соответствия курсов
course_mapping = {
    'Курс по электробезопасности': 'Курс по электробезопасности',
    'Безопасность и Охрана труда': 'Безопасность и Охрана труда',
    'Пожарная безопасность': 'Пожарная безопасность в объеме пожарно технического минимума',
    'Промышленная безопасность': 'Промышленная безопасность РК по производству работ на опасных производственных объектах'
}


class CertificateGenerator:
    def __init__(self):
        self.db = sqlite3.connect('certificates.db', check_same_thread=False, timeout=30)
        self.db.execute("PRAGMA journal_mode=WAL")
        self.create_tables()
        self.template_paths = {
            'Курс по электробезопасности': {
                'expert': Path('Элбез Expert Solution.docx'),
                'regular': Path('Элбез 2.docx')
            },
            'Безопасность и Охрана труда': Path('templates2.docx'),
            'Пожарная безопасность в объеме пожарно технического минимума': {
                'Для ответственных': Path('Птм ответсвенный.docx'),
                'Для работников': Path('Птм рабочие.docx')
            },
            'Промышленная безопасность РК по производству работ на опасных производственных объектах': Path(
                'templates3.docx')
        }
        self.output_dir = Path('generated_docs')
        self.output_dir.mkdir(exist_ok=True)

    def create_tables(self):
        self.db.execute('''CREATE TABLE IF NOT EXISTS certificates (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            protocol_number TEXT,
            fullname TEXT,
            workplace TEXT,
            job_title TEXT,
            position TEXT,
            group_number TEXT,
            cert_date DATE,
            next_date DATE,
            template_type TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )''')
        self.db.commit()

    def generate_document(self, data):
        try:
            generated_files = []
            protocol_number = str(self.db.execute("SELECT COUNT(*) FROM certificates").fetchone()[0] + 1).zfill(3)

            if data['position'] == 'Курс по электробезопасности':
                # Генерируем оба сертификата для электробезопасности
                group = data['qualification_group']
                if isinstance(group, str):
                    group = int(group) if group.isdigit() else 0
                replacements = {
                    '{{ input_fullname }}': data['fullname'],
                    '{{ input_place }}': data['workplace'],
                    '{{ input_jobtitle }}': data['job_title'],
                    '{{ input_what }}': "и выше" if group in [4, 5] else "",
                    '{{ input_do }}': "до " if group in [4, 5] else "",
                    '{{ input_nextdate }}': data['next_date'].strftime('%d.%m.%Y') if data['next_date'] else "",
                    '{{ input_proto }}': protocol_number,
                    '{{ input_date }}': data['cert_date'].strftime('%d.%m.%Y'),
                    '{{ input_group }}': str(group)
                }

                # Первый сертификат - Expert Solution
                doc1 = Document(self.template_paths['Курс по электробезопасности']['expert'])
                self._replace_text_in_document(doc1, replacements)
                output_path1 = self._save_document(doc1, f"{data['fullname']}_Expert")
                generated_files.append(output_path1)

                # Второй сертификат - Элбез 2
                doc2 = Document(self.template_paths['Курс по электробезопасности']['regular'])
                self._replace_text_in_document(doc2, replacements)
                output_path2 = self._save_document(doc2, f"{data['fullname']}_2")
                generated_files.append(output_path2)

            else:
                # Определяем путь к шаблону для других курсов
                if data['position'] == 'Пожарная безопасность в объеме пожарно технического минимума':
                    template_path = self.template_paths[data['position']][data['template_type']]
                else:
                    template_path = self.template_paths[data['position']]

                # Формируем замены в зависимости от типа курса
                if data['position'] == 'Безопасность и Охрана труда':
                    replacements = {
                        '{{ input_data }}': data['cert_date'].strftime('%d.%m.%Y'),
                        '{{ input_fullname }}': data['fullname'],
                        '{{ input_place }}': data['workplace'],
                        '{{ input_job }}': data['job_title'],
                        '{{ input_check }}': data['qualification_group'],
                        '{{ input_proto }}': protocol_number
                    }
                elif data[
                    'position'] == 'Промышленная безопасность РК по производству работ на опасных производственных объектах':
                    replacements = {
                        '{{ input_date }}': data['cert_date'].strftime('%d.%m.%Y'),
                        '{{ input_name }}': data['fullname'],
                        '{{ input_place }}': data['workplace'],
                        '{{ input_job }}': data['job_title'],
                        '{{ input_proto }}': protocol_number
                    }
                else:  # Пожарная безопасность
                    replacements = {
                        '{{ input_name }}': data['fullname'],
                        '{{ input_place }}': data['workplace'],
                        '{{ input_jobtitle }}': data['job_title'],
                        '{{ input_nextdate }}': data['next_date'].strftime('%d.%m.%Y') if data['next_date'] else "",
                        '{{ input_proto }}': protocol_number,
                        '{{ input_date }}': data['cert_date'].strftime('%d.%m.%Y')
                    }

                doc = Document(template_path)
                self._replace_text_in_document(doc, replacements)
                output_path = self._save_document(doc, data['fullname'])
                generated_files.append(output_path)

            # Сохраняем в базу данных
            self._save_to_database(data, protocol_number)
            return True, generated_files

        except Exception as e:
            logger.error(f"Ошибка при генерации документа: {str(e)}")
            return False, None

    def _replace_text_in_document(self, doc, replacements):
        # Замена в параграфах
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, str(value))

        # Замена в таблицах
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in replacements.items():
                        if key in cell.text:
                            cell.text = cell.text.replace(key, str(value))

    def _save_document(self, doc, fullname):
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_path = self.output_dir / f"{fullname}_{timestamp}.docx"
        doc.save(output_path)
        return output_path

    def _save_to_database(self, data, protocol_number):
        try:
            self.db.execute('''INSERT INTO certificates 
                (protocol_number, fullname, workplace, position, group_number, 
                 cert_date, next_date, template_type, job_title)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)''',
                            (protocol_number,
                             data['fullname'],
                             data['workplace'],
                             data['position'],
                             str(data['qualification_group']),
                             data['cert_date'],
                             data['next_date'],
                             data.get('template_type', ''),
                             data['job_title']))
            self.db.commit()
        except Exception as e:
            logger.error(f"Ошибка сохранения в БД: {e}")
            # Продолжаем работу даже при ошибке БД
            pass


async def handle_start_variations(update: Update, context: ContextTypes.DEFAULT_TYPE):
    message_text = update.message.text.lower()

    if message_text in start_variations:
        context.user_data.clear()
        await update.message.reply_text(
            MESSAGES['start'],
            reply_markup=ReplyKeyboardRemove()
        )
        return PEOPLE_COUNT
    return ConversationHandler.END


async def get_people_count(update: Update, context: ContextTypes.DEFAULT_TYPE):
    try:
        count = int(update.message.text.strip())
        if count <= 0:
            raise ValueError()

        context.user_data.clear()
        context.user_data.update({
            'total_people': count,
            'names': [],
            'workplaces': [],
            'job_titles': [],
            'positions': [],
            'groups': [],
            'cert_dates': [],
            'next_dates': [],
            'fire_safety_types': [],
            'check_types': []
        })

        if count == 1:
            await update.message.reply_text(
                "👤 Введите ФИО для 1 человека\n\n"
                "Пример:\n"
                "Иванов Андрей Андреевич"
            )
        else:
            await update.message.reply_text(
                MESSAGES['enter_names'].format(count) + "\n\n"
                                                        "Пример:\n" + "\n".join(
                    [f"Иванов Иван Иванович" for _ in range(min(count, 3))])
            )
        return FULLNAME

    except ValueError:
        await update.message.reply_text(MESSAGES['invalid_count'])
        return PEOPLE_COUNT


async def get_name(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = update.message.text.strip()
    total_people = context.user_data['total_people']
    names = [name.strip() for name in text.split('\n') if name.strip()]

    if len(names) != total_people:
        if total_people == 1:
            await update.message.reply_text(
                "❌ Необходимо ввести ФИО для 1 человека\n\n"
                "Пример:\n"
                "Иванов Андрей Андреевич"
            )
        else:
            await update.message.reply_text(
                f"❌ Необходимо ввести {total_people} ФИО\n"
                f"Вы ввели {len(names)} ФИО\n\n"
                f"Введите {total_people} ФИО, каждое с новой строки:\n\n"
                "Пример:\n" + "\n".join([f"Иванов Иван Иванович" for _ in range(min(total_people, 3))])
            )
        return FULLNAME

    invalid_names = []
    for i, name in enumerate(names, 1):
        name_parts = name.split()
        if len(name_parts) < 2 or not all(part.replace('-', '').isalpha() for part in name_parts):
            invalid_names.append(f"{i}) {name}")

    if invalid_names:
        await update.message.reply_text(
            "❌ Следующие ФИО введены неверно:\n" +
            "\n".join(invalid_names) + "\n\n"
                                       "Требования к ФИО:\n"
                                       "- Минимум 2 слова\n"
                                       "- Только буквы и дефис\n"
                                       "- Каждое ФИО с новой строки\n\n"
                                       f"Введите все {total_people} ФИО заново"
        )
        return FULLNAME

    context.user_data['names'] = names
    await update.message.reply_text(
        f"🏢 Введите {'название организации' if total_people == 1 else 'названия организаций для ' + str(total_people) + ' человек'}\n\n"
        "Пример:\n" + "\n".join([f"ТОО \"Компания {i + 1}\"" for i in range(min(total_people, 3))])
    )
    return WORKPLACE


async def get_workplace(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = update.message.text.strip()
    total_people = context.user_data['total_people']
    workplaces = [wp.strip() for wp in text.split('\n') if wp.strip()]

    if len(workplaces) != total_people:
        await update.message.reply_text(
            f"❌ Необходимо ввести {total_people} {'организацию' if total_people == 1 else 'организаций'}\n"
            f"Вы ввели {len(workplaces)} {'организацию' if len(workplaces) == 1 else 'организаций'}\n\n"
            f"Введите {total_people} {'организацию' if total_people == 1 else 'организаций'}, "
            f"{'каждую ' if total_people > 1 else ''}с новой строки:\n\n"
            "Пример:\n" + "\n".join([f"ТОО \"Компания {i + 1}\"" for i in range(min(total_people, 3))])
        )
        return WORKPLACE

    context.user_data['workplaces'] = workplaces
    await update.message.reply_text(
        f"👔 Введите {'должность' if total_people == 1 else 'должности для ' + str(total_people) + ' человек'}:\n\n"
        "Пример:\n" + "\n".join(["Инженер-энергетик", "Электрик", "Начальник участка"][:min(total_people, 3)])
    )
    return JOB_TITLE


async def get_job_title(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = update.message.text.strip()
    total_people = context.user_data['total_people']
    job_titles = [job.strip() for job in text.split('\n') if job.strip()]

    if len(job_titles) != total_people:
        if total_people == 1:
            await update.message.reply_text(
                "❌ Необходимо ввести должность\n\n"
                "Введите должность:\n\n"
                "Пример:\n"
                "Инженер-энергетик"
            )
        else:
            await update.message.reply_text(
                f"❌ Необходимо ввести {total_people} должностей\n"
                f"Вы ввели {len(job_titles)} должностей\n\n"
                f"Введите {total_people} должностей, каждую с новой строки:\n\n"
                "Пример:\n" + "\n".join(["Инженер-энергетик", "Электрик", "Начальник участка"][:min(total_people, 3)])
            )
        return JOB_TITLE

    context.user_data['job_titles'] = job_titles
    keyboard = [[pos] for pos in POSITIONS]
    reply_markup = ReplyKeyboardMarkup(keyboard, one_time_keyboard=True)

    await update.message.reply_text(
        f"📚 Выберите курс обучения для {context.user_data['names'][len(context.user_data['positions'])]}:",
        reply_markup=reply_markup
    )
    return POSITION


async def get_position(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = update.message.text.strip()

    # Инициализируем списки, если они еще не существуют
    if 'positions' not in context.user_data:
        context.user_data['positions'] = []
    if 'names' not in context.user_data:
        return ConversationHandler.END

    current_index = len(context.user_data['positions'])

    # Проверка на выход за пределы списка
    if current_index >= len(context.user_data['names']):
        await update.message.reply_text("❌ Произошла ошибка. Начните заново с команды /start")
        return ConversationHandler.END

    # Если это первый человек или предыдущий курс отличается
    if current_index == 0 or (current_index > 0 and context.user_data['positions'][-1] != text):
        selected_course = course_mapping.get(text, text)
    else:
        # Автоматически используем тот же курс, что и для предыдущего человека
        selected_course = context.user_data['positions'][-1]

    if selected_course == 'Курс по электробезопасности':
        context.user_data['positions'].append(selected_course)
        keyboard = [['2', '3', '4', '5']]
        reply_markup = ReplyKeyboardMarkup(keyboard, one_time_keyboard=True)
        await update.message.reply_text(
            f"Выберите группу допуска для {context.user_data['names'][current_index]}:",
            reply_markup=reply_markup
        )
        return GROUP

    elif selected_course == 'Пожарная безопасность в объеме пожарно технического минимума':
        context.user_data['positions'].append(selected_course)
        # Если предыдущий человек выбирал этот же курс, используем тот же тип
        if current_index > 0 and context.user_data['positions'][current_index - 1] == selected_course:
            prev_type = context.user_data['fire_safety_types'][-1]
            context.user_data['fire_safety_types'].append(prev_type)
            context.user_data['groups'].append('')

            if current_index + 1 < context.user_data['total_people']:
                # Спрашиваем, хочет ли пользователь использовать тот же курс для следующего человека
                keyboard = [
                    [f'✅ Тот же курс для {context.user_data["names"][current_index + 1]}'],
                    ['❌ Выбрать другой курс']
                ]
                reply_markup = ReplyKeyboardMarkup(keyboard, one_time_keyboard=True)
                await update.message.reply_text(
                    f"Использовать курс «{selected_course}» ({prev_type}) для {context.user_data['names'][current_index + 1]}?",
                    reply_markup=reply_markup
                )
                return POSITION
            else:
                await update.message.reply_text(
                    "📅 Введите дату выдачи (дд.мм.гггг):\n"
                    "Например: 01.01.2024"
                )
                return CERT_DATE
        else:
            keyboard = [['1. Для ответственных'], ['2. Для работников']]
            reply_markup = ReplyKeyboardMarkup(keyboard, one_time_keyboard=True)
            await update.message.reply_text(
                f"Выберите тип обучения для {context.user_data['names'][current_index]}:",
                reply_markup=reply_markup
            )
            return GROUP

    elif selected_course == 'Безопасность и Охрана труда':
        context.user_data['positions'].append(selected_course)
        context.user_data['groups'].append('')
        # Если предыдущий человек выбирал этот же курс, используем тот же тип проверки
        if current_index > 0 and context.user_data['positions'][current_index - 1] == selected_course:
            prev_type = context.user_data['check_types'][-1]
            context.user_data['check_types'].append(prev_type)

            if current_index + 1 < context.user_data['total_people']:
                # Спрашиваем, хочет ли пользователь использовать тот же курс для следующего человека
                keyboard = [
                    [f'✅ Тот же курс для {context.user_data["names"][current_index + 1]}'],
                    ['❌ Выбрать другой курс']
                ]
                reply_markup = ReplyKeyboardMarkup(keyboard, one_time_keyboard=True)
                await update.message.reply_text(
                    f"Использовать курс «{selected_course}» ({prev_type} проверка) для {context.user_data['names'][current_index + 1]}?",
                    reply_markup=reply_markup
                )
                return POSITION
            else:
                await update.message.reply_text(
                    "📅 Введите дату выдачи (дд.мм.гггг):\n"
                    "Например: 01.01.2024"
                )
                return CERT_DATE
        else:
            keyboard = [['первичный'], ['периодический']]
            reply_markup = ReplyKeyboardMarkup(keyboard, one_time_keyboard=True)
            await update.message.reply_text(
                f"Выберите тип проверки для {context.user_data['names'][current_index]}:",
                reply_markup=reply_markup
            )
            return CHECK_TYPE

    elif '✅ Тот же курс для' in text:
        # Автоматически используем тот же курс и настройки, что и для предыдущего человека
        prev_course = context.user_data['positions'][-1]
        context.user_data['positions'].append(prev_course)

        if prev_course == 'Пожарная безопасность в объеме пожарно технического минимума':
            prev_type = context.user_data['fire_safety_types'][-1]
            context.user_data['fire_safety_types'].append(prev_type)
            context.user_data['groups'].append('')
        elif prev_course == 'Безопасность и Охрана труда':
            prev_type = context.user_data['check_types'][-1]
            context.user_data['check_types'].append(prev_type)
            context.user_data['groups'].append('')
        elif prev_course == 'Курс по электробезопасности':
            prev_group = context.user_data['groups'][-1]
            context.user_data['groups'].append(prev_group)
        else:
            context.user_data['groups'].append('')

        if current_index + 1 < context.user_data['total_people']:
            # Спрашиваем про следующего человека
            keyboard = [
                [f'✅ Тот же курс для {context.user_data["names"][current_index + 1]}'],
                ['❌ Выбрать другой курс']
            ]
            reply_markup = ReplyKeyboardMarkup(keyboard, one_time_keyboard=True)

            additional_info = ""
            if prev_course == 'Курс по электробезопасности':
                additional_info = f" (группа {prev_group})"
            elif prev_course == 'Пожарная безопасность в объеме пожарно технического минимума':
                additional_info = f" ({prev_type})"
            elif prev_course == 'Безопасность и Охрана труда':
                additional_info = f" ({prev_type} проверка)"

            await update.message.reply_text(
                f"Использовать курс «{prev_course}»{additional_info} для {context.user_data['names'][current_index + 1]}?",
                reply_markup=reply_markup
            )
            return POSITION
        else:
            await update.message.reply_text(
                "📅 Введите дату выдачи (дд.мм.гггг):\n"
                "Например: 01.01.2024"
            )
            return CERT_DATE

    elif '❌ Выбрать другой курс' in text:
        keyboard = [[pos] for pos in POSITIONS]
        reply_markup = ReplyKeyboardMarkup(keyboard, one_time_keyboard=True)
        await update.message.reply_text(
            f"📚 Выберите курс обучения для {context.user_data['names'][current_index]}:",
            reply_markup=reply_markup
        )
        return POSITION

    elif selected_course in course_mapping.values():
        context.user_data['positions'].append(selected_course)
        context.user_data['groups'].append('')

        if current_index + 1 < context.user_data['total_people']:
            # Спрашиваем, хочет ли пользователь использовать тот же курс для следующего человека
            keyboard = [
                [f'✅ Тот же курс для {context.user_data["names"][current_index + 1]}'],
                ['❌ Выбрать другой курс']
            ]
            reply_markup = ReplyKeyboardMarkup(keyboard, one_time_keyboard=True)
            await update.message.reply_text(
                f"Использовать курс «{selected_course}» для {context.user_data['names'][current_index + 1]}?",
                reply_markup=reply_markup
            )
            return POSITION
        else:
            await update.message.reply_text(
                "📅 Введите дату выдачи (дд.мм.гггг):\n"
                "Например: 01.01.2024"
            )
            return CERT_DATE
    else:
        keyboard = [[pos] for pos in POSITIONS]
        reply_markup = ReplyKeyboardMarkup(keyboard, one_time_keyboard=True)
        await update.message.reply_text(
            "❌ Пожалуйста, выберите курс из списка",
            reply_markup=reply_markup
        )
        return POSITION


async def get_group(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = update.message.text.strip()
    current_index = len(context.user_data.get('groups', []))
    current_name = context.user_data['names'][current_index]
    position = context.user_data['positions'][current_index]

    if position == 'Курс по электробезопасности':
        try:
            group = int(text)
            if 2 <= group <= 5:
                if 'groups' not in context.user_data:
                    context.user_data['groups'] = []
                context.user_data['groups'].append(group)

                # Если есть ещё люди
                if current_index + 1 < context.user_data['total_people']:
                    keyboard = [[pos] for pos in POSITIONS]
                    reply_markup = ReplyKeyboardMarkup(keyboard, one_time_keyboard=True)
                    await update.message.reply_text(
                        f"Выберите курс для {context.user_data['names'][current_index + 1]}:",
                        reply_markup=reply_markup
                    )
                    return POSITION
                else:
                    await update.message.reply_text(
                        "Введите дату выдачи (дд.мм.гггг):\n"
                        "Например: 01.01.2024"
                    )
                    return CERT_DATE
            else:
                keyboard = [['2', '3', '4', '5']]
                reply_markup = ReplyKeyboardMarkup(keyboard, one_time_keyboard=True)
                await update.message.reply_text(
                    f"❌ Выберите группу допуска от 2 до 5 для {current_name}:",
                    reply_markup=reply_markup
                )
                return GROUP
        except ValueError:
            keyboard = [['2', '3', '4', '5']]
            reply_markup = ReplyKeyboardMarkup(keyboard, one_time_keyboard=True)
            await update.message.reply_text(
                f"❌ Выберите группу допуска от 2 до 5 для {current_name}:",
                reply_markup=reply_markup
            )
            return GROUP

    elif position == 'Пожарная безопасность в объеме пожарно технического минимума':
        if text.startswith('1.'):
            if 'fire_safety_types' not in context.user_data:
                context.user_data['fire_safety_types'] = []
            context.user_data['fire_safety_types'].append('Для ответственных')
            context.user_data.setdefault('groups', []).append('')
        elif text.startswith('2.'):
            if 'fire_safety_types' not in context.user_data:
                context.user_data['fire_safety_types'] = []
            context.user_data['fire_safety_types'].append('Для работников')
            context.user_data.setdefault('groups', []).append('')
        else:
            keyboard = [['1. Для ответственных'], ['2. Для работников']]
            reply_markup = ReplyKeyboardMarkup(keyboard, one_time_keyboard=True)
            await update.message.reply_text(
                f"❌ Выберите тип обучения для {current_name}:",
                reply_markup=reply_markup
            )
            return GROUP

        # Если есть ещё люди
        if current_index + 1 < context.user_data['total_people']:
            keyboard = [[pos] for pos in POSITIONS]
            reply_markup = ReplyKeyboardMarkup(keyboard, one_time_keyboard=True)
            await update.message.reply_text(
                f"Выберите курс для {context.user_data['names'][current_index + 1]}:",
                reply_markup=reply_markup
            )
            return POSITION
        else:
            await update.message.reply_text(
                "Введите дату выдачи (дд.мм.гггг):\n"
                "Например: 01.01.2024"
            )
            return CERT_DATE

    else:
        # Для остальных курсов
        context.user_data.setdefault('groups', []).append('')
        if current_index + 1 < context.user_data['total_people']:
            keyboard = [[pos] for pos in POSITIONS]
            reply_markup = ReplyKeyboardMarkup(keyboard, one_time_keyboard=True)
            await update.message.reply_text(
                f"Выберите курс для {context.user_data['names'][current_index + 1]}:",
                reply_markup=reply_markup
            )
            return POSITION
        else:
            await update.message.reply_text(
                "Введите дату выдачи (дд.мм.гггг):\n"
                "Например: 01.01.2024"
            )
            return CERT_DATE


async def get_cert_date(update: Update, context: ContextTypes.DEFAULT_TYPE):
    try:
        if 'total_people' not in context.user_data:
            logger.error("total_people отсутствует в контексте")
            await update.message.reply_text(
                "❌ Произошла ошибка. Пожалуйста, начните заново с команды /start"
            )
            return ConversationHandler.END

        # Инициализация списков, если их нет
        if 'next_dates' not in context.user_data:
            context.user_data['next_dates'] = [None] * context.user_data['total_people']
        if 'cert_dates' not in context.user_data:
            context.user_data['cert_dates'] = []

        text = update.message.text.strip()
        total_people = context.user_data['total_people']

        # Если это первый вызов функции
        if not text or text in ['/start', 'создать', '✅ Создать сертификаты', '❌ Начать заново'] or any(
                pos in text for pos in POSITIONS):
            message = f"📅 Введите дату выдачи (дд.мм.гггг) для {total_people} {'человека' if 2 <= total_people % 10 <= 4 and (total_people < 10 or total_people > 20) else 'человек'}:\n\n"

            # Добавляем информацию о каждом человеке
            for idx, name in enumerate(context.user_data['names'], 1):
                position = context.user_data['positions'][idx - 1]
                message += f"{idx}) {name}\n"
                message += f"    Курс: {position}\n"
                if position == 'Курс по электробезопасности':
                    group = context.user_data.get('groups', [])[idx - 1]
                    message += f"    Группа: {group}\n"
                message += "\n"

            message += "Пример ввода дат:\n"
            example_dates = []
            for _ in range(total_people):
                example_dates.append("03.12.2024")
            message += "\n".join(example_dates)
            await update.message.reply_text(message)
            return CERT_DATE

        dates = [date.strip() for date in text.split('\n') if date.strip()]

        # Проверяем количество дат
        if len(dates) != total_people:
            message = f"❌ Необходимо ввести {total_people} {'дату' if total_people == 1 else 'даты' if 2 <= total_people % 10 <= 4 and (total_people < 10 or total_people > 20) else 'дат'}\n\n"
            message += f"📅 Введите дату выдачи (дд.мм.гггг) для {total_people} {'человека' if 2 <= total_people % 10 <= 4 and (total_people < 10 or total_people > 20) else 'человек'}:\n\n"

            # Добавляем информацию о каждом человеке
            for idx, name in enumerate(context.user_data['names'], 1):
                position = context.user_data['positions'][idx - 1]
                message += f"{idx}) {name}\n"
                message += f"    Курс: {position}\n"
                if position == 'Курс по электробезопасности':
                    group = context.user_data.get('groups', [])[idx - 1]
                    message += f"    Группа: {group}\n"
                message += "\n"

            message += "Пример ввода дат:\n"
            example_dates = []
            for _ in range(total_people):
                example_dates.append("03.12.2024")
            message += "\n".join(example_dates)
            await update.message.reply_text(message)
            return CERT_DATE

        # Проверяем форматы дат
        cert_dates = []
        try:
            for date_str in dates:
                cert_date = datetime.strptime(date_str, '%d.%m.%Y')
                if cert_date > datetime.now():
                    message = "❌ Дата не может быть в будущем\n\n"
                    message += f"📅 Введите дату выдачи (дд.мм.гггг) для {total_people} {'человека' if 2 <= total_people % 10 <= 4 and (total_people < 10 or total_people > 20) else 'человек'}:\n\n"

                    # Добавляем информацию о каждом человеке
                    for idx, name in enumerate(context.user_data['names'], 1):
                        position = context.user_data['positions'][idx - 1]
                        message += f"{idx}) {name}\n"
                        message += f"    Курс: {position}\n"
                        if position == 'Курс по электробезопасности':
                            group = context.user_data.get('groups', [])[idx - 1]
                            message += f"    Группа: {group}\n"
                        message += "\n"

                    message += "Пример ввода дат:\n"
                    message += "\n".join(["01.01.2024" for _ in range(total_people)])
                    await update.message.reply_text(message)
                    return CERT_DATE
                cert_dates.append(cert_date)
        except ValueError:
            message = "❌ Неверный формат даты\n\n"
            message += f"📅 Введите дату выдачи (дд.мм.гггг) для {total_people} {'человека' if 2 <= total_people % 10 <= 4 and (total_people < 10 or total_people > 20) else 'человек'}:\n\n"

            # Добавляем информацию о каждом человеке
            for idx, name in enumerate(context.user_data['names'], 1):
                position = context.user_data['positions'][idx - 1]
                message += f"{idx}) {name}\n"
                message += f"    Курс: {position}\n"
                if position == 'Курс по электробезопасности':
                    group = context.user_data.get('groups', [])[idx - 1]
                    message += f"    Группа: {group}\n"
                message += "\n"

            message += "Пример ввода дат:\n"
            message += "\n".join(["01.01.2024" for _ in range(total_people)])
            await update.message.reply_text(message)
            return CERT_DATE

        context.user_data['cert_dates'] = cert_dates

        # Определяем, для кого нужны следующие даты
        needs_next_dates = []
        for i, pos in enumerate(context.user_data['positions']):
            if pos in ['Курс по электробезопасности', 'Пожарная безопасность в объеме пожарно технического минимума']:
                needs_next_dates.append(i)

        if needs_next_dates:
            message = f"📅 Введите дату следующей проверки (дд.мм.гггг) для {len(needs_next_dates)} {'человека' if 2 <= len(needs_next_dates) % 10 <= 4 and (len(needs_next_dates) < 10 or len(needs_next_dates) > 20) else 'человек'}:\n\n"

            # Добавляем информацию о каждом человеке
            for idx, i in enumerate(needs_next_dates, 1):
                name = context.user_data['names'][i]
                position = context.user_data['positions'][i]
                cert_date = cert_dates[i].strftime('%d.%m.%Y')

                message += f"{idx}) {name}\n"
                message += f"    Курс: {position}\n"
                message += f"    Дата выдачи: {cert_date}\n"
                if position == 'Курс по электробезопасности':
                    group = context.user_data.get('groups', [])[i]
                    message += f"    Группа: {group}\n"
                message += "\n"

            message += "Пример ввода дат:\n"
            example_dates = []
            for _ in range(len(needs_next_dates)):
                example_dates.append("03.12.2025")
            message += "\n".join(example_dates)

            await update.message.reply_text(message)
            context.user_data['needs_next_dates'] = needs_next_dates
            return NEXT_DATE
        else:
            context.user_data['next_dates'] = [None] * total_people
            return await show_summary(update, context)

    except Exception as e:
        logger.error(f"Ошибка обработки дат: {e}")
        await update.message.reply_text(
            "❌ Произошла ошибка. Введите даты заново"
        )
        return CERT_DATE


async def get_next_date(update: Update, context: ContextTypes.DEFAULT_TYPE):
    try:
        text = update.message.text.strip()
        needs_next_dates = context.user_data.get('needs_next_dates', [])

        # Проверяем, есть ли needs_next_dates в контексте
        if not needs_next_dates:
            logger.error("needs_next_dates отсутствует в контексте")
            await update.message.reply_text(
                "❌ Произошла ошибка. Пожалуйста, начните заново с команды /start"
            )
            return ConversationHandler.END

        # Если это первый вызов функции или текст содержит даты
        if all(str(i) in text for i in range(10)):  # Проверяем наличие цифр в тексте
            message = f"📅 Введите дату следующей проверки (дд.мм.гггг) для {len(needs_next_dates)} {'человека' if 2 <= len(needs_next_dates) % 10 <= 4 and (len(needs_next_dates) < 10 or len(needs_next_dates) > 20) else 'человек'}:\n\n"

            # Проверяем наличие всех необходимых данных
            if not all(key in context.user_data for key in ['names', 'positions', 'cert_dates']):
                logger.error("Отсутствуют необходимые данные в контексте")
                await update.message.reply_text(
                    "❌ Произошла ошибка. Пожалуйста, начните заново с команды /start"
                )
                return ConversationHandler.END

            # Добавляем информацию о каждом человеке
            for idx, i in enumerate(needs_next_dates, 1):
                try:
                    name = context.user_data['names'][i]
                    position = context.user_data['positions'][i]
                    cert_date = context.user_data['cert_dates'][i].strftime('%d.%m.%Y')

                    message += f"{idx}) {name}\n"
                    message += f"    Курс: {position}\n"
                    message += f"    Дата выдачи: {cert_date}\n"
                    if position == 'Курс по электробезопасности':
                        group = context.user_data.get('groups', [])[i]
                        message += f"    Группа: {group}\n"
                    message += "\n"
                except IndexError as e:
                    logger.error(f"Ошибка индексации при формировании сообщения: {e}")
                    await update.message.reply_text(
                        "❌ Произошла ошибка. Пожалуйста, начните заново с команды /start"
                    )
                    return ConversationHandler.END

            message += "Пример ввода дат:\n"
            message += "\n".join(["03.12.2025" for _ in range(len(needs_next_dates))])

            await update.message.reply_text(message)
            return NEXT_DATE

        # Обработка введенных дат
        dates = [date.strip() for date in text.split('\n') if date.strip()]

        if len(dates) != len(needs_next_dates):
            message = f"❌ Необходимо ввести {len(needs_next_dates)} {'дату' if len(needs_next_dates) == 1 else 'даты' if 2 <= len(needs_next_dates) % 10 <= 4 and (len(needs_next_dates) < 10 or len(needs_next_dates) > 20) else 'дат'}\n\n"
            message += f"Введите даты заново:\n"
            message += "\n".join(["01.02.2024" for _ in range(len(needs_next_dates))])
            await update.message.reply_text(message)
            return NEXT_DATE

        # Инициализация списка следующих дат
        next_dates = [None] * context.user_data['total_people']
        invalid_dates = []

        # Проверка и обработка дат
        for i, date_str in enumerate(dates):
            try:
                next_date = datetime.strptime(date_str, '%d.%m.%Y')
                person_index = needs_next_dates[i]

                if person_index >= len(context.user_data['cert_dates']):
                    logger.error(f"Индекс {person_index} выходит за пределы списка cert_dates")
                    raise IndexError("Индекс выходит за пределы списка")

                cert_date = context.user_data['cert_dates'][person_index]

                if next_date <= cert_date:
                    name = context.user_data['names'][person_index]
                    invalid_dates.append(
                        f"❌ {name}:\n"
                        f"    Дата следующей проверки ({next_date.strftime('%d.%m.%Y')})\n"
                        f"    должна быть позже даты выдачи ({cert_date.strftime('%d.%m.%Y')})"
                    )
                else:
                    next_dates[person_index] = next_date

            except (ValueError, IndexError) as e:
                logger.error(f"Ошибка при обработке даты: {e}")
                await update.message.reply_text(
                    "❌ Произошла ошибка при обработке дат. Пожалуйста, введите даты заново в формате дд.мм.гггг"
                )
                return NEXT_DATE

        if invalid_dates:
            message = "❌ Обнаружены ошибки:\n\n"
            message += "\n\n".join(invalid_dates)
            message += "\n\nВведите все даты заново:\n"
            message += "\n".join(["01.02.2024" for _ in range(len(needs_next_dates))])
            await update.message.reply_text(message)
            return NEXT_DATE

        context.user_data['next_dates'] = next_dates
        return await show_summary(update, context)

    except Exception as e:
        logger.error(f"Ошибка обработки следующих дат: {e}")
        await update.message.reply_text(
            "❌ Произошла ошибка. Пожалуйста, начните заново с команды /start"
        )
        return ConversationHandler.END


async def show_summary(update: Update, context: ContextTypes.DEFAULT_TYPE):
    summary = await create_summary(context)
    keyboard = [['✅ Создать сертификаты', '❌ Начать заново']]
    reply_markup = ReplyKeyboardMarkup(keyboard, one_time_keyboard=True)
    await update.message.reply_text(
        f"Проверьте данные:\n\n{summary}",
        reply_markup=reply_markup
    )
    return CONFIRM_DATA


async def create_summary(context) -> str:
    # Создаем заголовок
    summary = "╔══════════════════════════╗\n"
    summary += "║    ВЫДАЧА УДОСТОВЕРЕНИЙ    ║\n"
    summary += "╚══════════════════════════╝\n\n"
    summary += "📋 Итоговая сводка:\n\n"

    # Для каждого человека
    for i in range(context.user_data['total_people']):
        # Создаем разделитель для каждого человека
        summary += "┌" + "─" * 40 + "┐\n"
        summary += f"│ 👤 Человек {i + 1}" + " " * (29 - len(str(i + 1))) + "│\n"
        summary += "├" + "─" * 40 + "┤\n"

        # Основная информация
        name_line = f"│ ФИО: {context.user_data['names'][i]}"
        summary += name_line + " " * (41 - len(name_line)) + "│\n"

        workplace_line = f"│ Организация: {context.user_data['workplaces'][i]}"
        summary += workplace_line + " " * (41 - len(workplace_line)) + "│\n"

        job_line = f"│ Должность: {context.user_data['job_titles'][i]}"
        summary += job_line + " " * (41 - len(job_line)) + "│\n"

        position_line = f"│ Курс: {context.user_data['positions'][i]}"
        summary += position_line + " " * (41 - len(position_line)) + "│\n"

        # Дополнительная информация для ПТМ
        if context.user_data['positions'][i] == 'Пожарная безопасность в объеме пожарно технического минимума':
            type_line = f"│ Тип: {context.user_data['fire_safety_types'][i]}"
            summary += type_line + " " * (41 - len(type_line)) + "│\n"

        # Информация о группе
        if context.user_data['positions'][i] == 'Курс по электробезопасности':
            group_line = f"│ Группа: {context.user_data['groups'][i]}"
            summary += group_line + " " * (41 - len(group_line)) + "│\n"
        elif context.user_data['positions'][i] == 'Безопасность и Охрана труда':
            check_type_line = f"│ Тип проверки: {context.user_data['check_types'][i]}"
            summary += check_type_line + " " * (41 - len(check_type_line)) + "│\n"

        # Даты
        cert_date = context.user_data['cert_dates'][i].strftime('%d.%m.%Y')
        date_line = f"│ Дата выдачи: {cert_date}"
        summary += date_line + " " * (41 - len(date_line)) + "│\n"

        if context.user_data['next_dates'][i]:
            next_date = context.user_data['next_dates'][i].strftime('%d.%m.%Y')
            next_date_line = f"│ Следующая проверка: {next_date}"
            summary += next_date_line + " " * (41 - len(next_date_line)) + "│\n"

        # Закрываем рамку для текущего человека
        summary += "└" + "─" * 40 + "┘\n\n"

    # Добавляем вопрос в конце
    summary += "Создать документы?"

    return summary


async def confirm_data(update: Update, context: ContextTypes.DEFAULT_TYPE):
    response = update.message.text.strip().lower()

    if '✅' in response or 'да' in response:
        await update.message.reply_text(
            "⚙️ Начинаю генерацию документов...",
            reply_markup=ReplyKeyboardRemove()
        )

        generator = CertificateGenerator()

        for i in range(context.user_data['total_people']):
            try:
                data = {
                    'fullname': context.user_data['names'][i],
                    'workplace': context.user_data['workplaces'][i],
                    'job_title': context.user_data['job_titles'][i],
                    'position': context.user_data['positions'][i],
                    'qualification_group': context.user_data.get('groups', [''])[i] or '',
                    'cert_date': context.user_data['cert_dates'][i],
                    'next_date': context.user_data.get('next_dates', [None])[i],
                    'template_type': ''
                }

                if data['position'] == 'Пожарная безопасность в объеме пожарно технического минимума':
                    data['template_type'] = context.user_data.get('fire_safety_types', [])[i]
                elif data['position'] == 'Безопасность и Охрана труда':
                    data['qualification_group'] = context.user_data.get('check_types', [])[i]

                success, file_paths = generator.generate_document(data)

                if success and file_paths:
                    for file_path in file_paths:
                        try:
                            await update.message.reply_document(
                                document=open(file_path, 'rb'),
                                filename=file_path.name,
                                caption=f"✅ Сертификат для {data['fullname']} создан!"
                            )
                        except Exception as e:
                            logger.error(f"Ошибка отправки файла {file_path}: {e}")
                            await update.message.reply_text(
                                f"⚠️ Сертификат создан, но возникла ошибка при отправке: {file_path.name}"
                            )
                else:
                    await update.message.reply_text(
                        f"❌ Ошибка создания сертификата для {data['fullname']}"
                    )

            except Exception as e:
                logger.error(f"Ошибка генерации: {e}")
                await update.message.reply_text(
                    f"❌ Ошибка при создании сертификата для {context.user_data['names'][i]}"
                )
                continue

        await update.message.reply_text("✅ Готово! Для создания новых сертификатов нажмите /start")
        return ConversationHandler.END

    elif '❌' in response or 'нет' in response:
        await update.message.reply_text(
            "🔄 Начинаем заново.\n\nВведите количество человек:",
            reply_markup=ReplyKeyboardRemove()
        )
        context.user_data.clear()
        return PEOPLE_COUNT

    else:
        keyboard = [['✅ Да, создать документы', '❌ Нет, начать заново']]
        reply_markup = ReplyKeyboardMarkup(keyboard, one_time_keyboard=True)
        await update.message.reply_text(
            "❌ Выберите один из вариантов:",
            reply_markup=reply_markup
        )
        return CONFIRM_DATA


async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(MESSAGES['cancelled'], reply_markup=ReplyKeyboardRemove())
    return ConversationHandler.END


async def get_check_type(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = update.message.text.strip()
    current_index = len(context.user_data.get('check_types', []))
    current_name = context.user_data['names'][current_index]

    if text in ['первичный', 'периодический']:
        if 'check_types' not in context.user_data:
            context.user_data['check_types'] = []
        context.user_data['check_types'].append(text)

        # Если есть ещё люди
        if current_index + 1 < context.user_data['total_people']:
            keyboard = [[pos] for pos in POSITIONS]
            reply_markup = ReplyKeyboardMarkup(keyboard, one_time_keyboard=True)
            await update.message.reply_text(
                f"Выберите курс для {context.user_data['names'][current_index + 1]}:",
                reply_markup=reply_markup
            )
            return POSITION
        else:
            await update.message.reply_text(
                "Введите дату выдачи (дд.мм.гггг):\n"
                "Например: 01.01.2024"
            )
            return CERT_DATE
    else:
        keyboard = [['первичный'], ['периодический']]
        reply_markup = ReplyKeyboardMarkup(keyboard, one_time_keyboard=True)
        await update.message.reply_text(
            f"❌ Выберите тип проверки для {current_name}:",
            reply_markup=reply_markup
        )
        return CHECK_TYPE


def main():
    # Настройка логирования
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
        handlers=[
            logging.FileHandler('bot.log', encoding='utf-8'),
            logging.StreamHandler()
        ]
    )

    try:
        # Создание и настройка приложения
        application = Application.builder().token("7661532547:AAE8wCPk_wi1OxE8jpTpqTS7Akd6pUWtu4Q").build()

        # Добавление обработчиков
        conv_handler = ConversationHandler(
            entry_points=[
                CommandHandler('start', handle_start_variations),
                MessageHandler(filters.Regex(f"^({'|'.join(start_variations)})$"), handle_start_variations)
            ],
            states={
                PEOPLE_COUNT: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_people_count)],
                FULLNAME: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_name)],
                WORKPLACE: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_workplace)],
                JOB_TITLE: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_job_title)],
                POSITION: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_position)],
                GROUP: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_group)],
                CERT_DATE: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_cert_date)],
                NEXT_DATE: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_next_date)],
                CONFIRM_DATA: [MessageHandler(filters.TEXT & ~filters.COMMAND, confirm_data)],
                CHECK_TYPE: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_check_type)]
            },
            fallbacks=[
                CommandHandler('cancel', cancel),
                CommandHandler('start', handle_start_variations),
                MessageHandler(filters.Regex(f"^({'|'.join(start_variations)})$"), handle_start_variations)
            ],
            allow_reentry=True
        )

        application.add_handler(conv_handler)

        # Запуск бота с параметрами для предотвращения конфликтов
        print("✅ Бот запущен и готов к работе!")
        print("ℹ️ Нажмите Ctrl+C для остановки")
        application.run_polling(
            allowed_updates=Update.ALL_TYPES,
            drop_pending_updates=True,
            close_loop=True
        )

    except Exception as e:
        print(f"❌ Ошибка при запуске бота: {str(e)}")
        logging.error(f"Ошибка при запуске бота: {str(e)}", exc_info=True)
        sys.exit(1)


if __name__ == '__main__':
    main()